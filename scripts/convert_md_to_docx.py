"""convert_md_to_docx.py
Simple CLI utility to convert a Markdown file to a Word .docx document.

Usage:
    python convert_md_to_docx.py <input_markdown> <output_docx>

Dependencies:
    pip install markdown python-docx

This intentionally avoids external binaries like Pandoc to keep the
conversion lightweight and self-contained.
"""

import sys
from pathlib import Path

import markdown  # type: ignore
from docx import Document  # type: ignore
from bs4 import BeautifulSoup  # type: ignore


def md_to_html(md_text: str) -> str:
    """Convert Markdown text to HTML string using python-markdown."""
    return markdown.markdown(md_text, extensions=["tables", "fenced_code"])  # enable basic extras


def html_to_docx(html: str, doc: Document) -> None:
    """Very lightweight HTML -> docx converter supporting paragraphs and headings.

    For enterprise-grade conversion use pandoc; this handles basic structures
    sufficient for internal documentation.
    """
    soup = BeautifulSoup(html, "html.parser")
    for elem in soup.children:
        if elem.name == "h1":
            doc.add_heading(elem.get_text(), level=1)
        elif elem.name == "h2":
            doc.add_heading(elem.get_text(), level=2)
        elif elem.name == "h3":
            doc.add_heading(elem.get_text(), level=3)
        elif elem.name == "h4":
            doc.add_heading(elem.get_text(), level=4)
        elif elem.name == "h5":
            doc.add_heading(elem.get_text(), level=5)
        elif elem.name == "h6":
            doc.add_heading(elem.get_text(), level=6)
        elif elem.name == "p":
            doc.add_paragraph(elem.get_text())
        elif elem.name == "ul":
            for li in elem.find_all("li", recursive=False):
                doc.add_paragraph(li.get_text(), style="List Bullet")
        elif elem.name == "ol":
            for li in elem.find_all("li", recursive=False):
                doc.add_paragraph(li.get_text(), style="List Number")
        elif elem.name == "pre":
            # code block: keep monospace formatting
            code = elem.get_text()
            doc.add_paragraph(code, style="Quote")
        # ignore other tags for brevity


def convert(md_path: Path, docx_path: Path) -> None:
    md_text = md_path.read_text(encoding="utf-8")
    html = md_to_html(md_text)
    doc = Document()
    html_to_docx(html, doc)
    doc.save(docx_path)


if __name__ == "__main__":
    if len(sys.argv) != 3:
        print("Usage: python convert_md_to_docx.py <input.md> <output.docx>")
        sys.exit(1)

    input_md = Path(sys.argv[1])
    output_docx = Path(sys.argv[2])

    if not input_md.exists():
        print(f"Error: {input_md} does not exist.")
        sys.exit(1)

    convert(input_md, output_docx)
    print(f"Converted {input_md} -> {output_docx}")
